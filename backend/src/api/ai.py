from fastapi import APIRouter, UploadFile, File, HTTPException
from pydantic import BaseModel
from tempfile import NamedTemporaryFile
from faster_whisper import WhisperModel
import os
from ..services.llm import generate_report_from_transcript
from fastapi import Response
from fastapi.responses import StreamingResponse
from io import BytesIO
from pydantic import BaseModel
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK
from fastapi.responses import JSONResponse
from pydantic import BaseModel, EmailStr
from ..services.docx import render_report_docx
from ..services.mailer import send_mail_with_attachment
import tempfile

router = APIRouter()


class ReportEmailRequest(BaseModel):
    report_md: str
    author: str
    report_date: str

@router.post("/report-email")
async def report_email(payload: ReportEmailRequest):
    try:
        # 1. Génère un .docx temporaire
        tmp_path = tempfile.mktemp(suffix=".docx")
        render_report_docx(payload.report_md, payload.author, payload.report_date, tmp_path)

        # 2. Envoie par mail
        await send_mail_with_attachment(
            to=payload.author,  # <-- email utilisateur
            subject=f"Compte rendu du {payload.report_date}",
            body="Veuillez trouver ci-joint le compte rendu généré automatiquement.",
            file_path=tmp_path,
        )
        return JSONResponse({"ok": True, "message": "Email envoyé ✅"})
    except Exception as e:
        raise HTTPException(status_code=500, detail=str(e))
# ---------- Whisper local (déjà configuré précédemment) ----------
_model = None
def get_model():
    global _model
    if _model is not None:
        return _model
    model_dir = os.getenv("WHISPER_MODEL_DIR")
    model_size = os.getenv("WHISPER_MODEL_SIZE", "base")
    device     = os.getenv("WHISPER_DEVICE", "cpu")
    compute    = os.getenv("WHISPER_COMPUTE", "int8")
    if not model_dir:
        raise RuntimeError("WHISPER_MODEL_DIR non défini")
    local_path = os.path.join(model_dir, model_size)
    if not os.path.isdir(local_path):
        raise RuntimeError(f"Modèle introuvable: {local_path}")
    _model = WhisperModel(local_path, device=device, compute_type=compute, cpu_threads=4, num_workers=1)
    return _model

class TranscribeOut(BaseModel):
    text: str
    language: str | None = None

@router.post("/transcribe", response_model=TranscribeOut)
async def transcribe(file: UploadFile = File(...)):
    suffix = os.path.splitext(file.filename or "audio.webm")[-1] or ".webm"
    with NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
        tmp.write(await file.read())
        tmp_path = tmp.name
    try:
        model = get_model()
        lang = os.getenv("WHISPER_FORCE_LANG") or None
        segments, info = model.transcribe(tmp_path, language=lang, vad_filter=True, beam_size=1)
        text = "".join(s.text for s in segments).strip()
        return TranscribeOut(text=text, language=getattr(info, "language", None))
    finally:
        try: os.remove(tmp_path)
        except: pass

# ---------- Nouveau: génération de compte rendu ----------
class ReportIn(BaseModel):
    transcript: str
    # 👇 métadonnées optionnelles (arrivent du Dashboard)
    author: str | None = None
    author_email: str | None = None
    job_title: str | None = None
    site: str | None = None
    report_date: str | None = None  # "JJ/MM/AAAA"

class ReportOut(BaseModel):
    report_md: str
class ReportOut(BaseModel):
    report_md: str

@router.post("/report", response_model=ReportOut)
async def report(payload: ReportIn):
    if not payload.transcript.strip():
        raise HTTPException(400, detail="Transcript vide")

    md = await generate_report_from_transcript(
        payload.transcript,
        author=payload.author or "",
        author_email=payload.author_email or "",
        job_title=payload.job_title or "",
        site=payload.site or "",
        report_date=payload.report_date or "",
    )
    return ReportOut(report_md=md)
# ---------- Nouveau: en une seule requête ----------
class TranscribeReportOut(BaseModel):
    text: str
    report_md: str
    language: str | None = None

@router.post("/transcribe-and-report", response_model=TranscribeReportOut)
async def transcribe_and_report(file: UploadFile = File(...)):
    # 1) transcrire
    t = await transcribe(file)  # réutilise la logique ci-dessus
    # 2) générer compte rendu
    md = await generate_report_from_transcript(t.text)
    return TranscribeReportOut(text=t.text, language=t.language, report_md=md)

class ReportDocxIn(BaseModel):
    report_md: str
    author: str
    author_email: EmailStr | None = None  # pas utilisé pour le download, mais utile si tu veux l’afficher
    report_date: str
    job_title: str | None = None
    site: str | None = None

def _add_md_to_docx(doc: Document, md: str):
    """
    Convertisseur Markdown -> docx très simple adapté à notre gabarit :
    - # Titre, ## Sous-titres
    - Puces '- ...'
    - Cases à cocher '- [ ] ...' ou '- [x] ...'
    - Paragraphes
    """
    lines = md.replace("\r\n", "\n").split("\n")
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph()  # ligne vide
            continue

        # Titres
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            continue

        # Cases à cocher
        if line.lstrip().startswith("- [ ] "):
            p = doc.add_paragraph()
            run = p.add_run("☐ " + line.lstrip()[5:])
            run.font.size = Pt(11)
            continue
        if line.lstrip().startswith("- [x] ") or line.lstrip().startswith("- [X] "):
            p = doc.add_paragraph()
            run = p.add_run("☑ " + line.lstrip()[5:])
            run.font.size = Pt(11)
            continue

        # Puces
        if line.lstrip().startswith("- "):
            p = doc.add_paragraph(line.lstrip()[2:], style="List Bullet")
            continue

        # Paragraphe standard
        doc.add_paragraph(line)

def build_docx(report_md: str, author: str, report_date: str) -> bytes:
    doc = Document()

    # Page de garde courte
    title = doc.add_heading("Compte rendu", level=0)
    # métadonnées
    meta = doc.add_paragraph()
    meta.add_run(f"Date : {report_date}\n").bold = True
    meta.add_run(f"Rédigé par : {author}\n").bold = True

    # saut de ligne
    doc.add_paragraph().add_run().add_break(WD_BREAK.LINE)

    # Contenu Markdown (mise en forme simple)
    _add_md_to_docx(doc, report_md)

    # exporter en mémoire
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()

@router.post("/report-docx")
async def report_docx(payload: ReportDocxIn):
    try:
        data = build_docx(payload.report_md, payload.author, payload.report_date)
        filename = "compte-rendu.docx"
        return StreamingResponse(
            BytesIO(data),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={"Content-Disposition": f'attachment; filename="{filename}"'}
        )
    except Exception as e:
        raise HTTPException(500, detail=f"DOCX error: {e}")