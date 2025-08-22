from io import BytesIO
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_BREAK

def _add_md_to_docx(doc: Document, md: str):
    lines = md.replace("\r\n", "\n").split("\n")
    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            doc.add_paragraph()
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1); continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2); continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3); continue
        if line.lstrip().startswith("- [ ] "):
            p = doc.add_paragraph(); p.add_run("☐ " + line.lstrip()[5:]).font.size = Pt(11); continue
        if line.lstrip().startswith(("- [x] ", "- [X] ")):
            p = doc.add_paragraph(); p.add_run("☑ " + line.lstrip()[5:]).font.size = Pt(11); continue
        if line.lstrip().startswith("- "):
            doc.add_paragraph(line.lstrip()[2:], style="List Bullet"); continue
        doc.add_paragraph(line)

def render_report_docx(report_md: str, author: str, report_date: str, out_path: str | None = None) -> bytes:
    doc = Document()
    doc.add_heading("Compte rendu", level=0)
    meta = doc.add_paragraph()
    meta.add_run(f"Date : {report_date}\n").bold = True
    meta.add_run(f"Rédigé par : {author}\n").bold = True
    doc.add_paragraph().add_run().add_break(WD_BREAK.LINE)
    _add_md_to_docx(doc, report_md)

    if out_path:
        doc.save(out_path)
        return b""
    buf = BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()